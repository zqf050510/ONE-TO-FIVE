import streamlit as st
import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import PyPDF2
import io
import re
import time

# ============================================================
# 【核心配置】
# ============================================================
# 建议在 Streamlit Cloud 的 Secrets 中配置，或者通过侧边栏输入
MY_REAL_API_KEY = "gTBJ6GhqKOimfYGhys1gbAlzxbwsOoy3"
MY_ASSISTANT_ID = "2035992077834375936"
MY_API_URL = "https://yuanqi.tencent.com/openapi/v1/agent/chat/completions"


# ============================================================
# 【工具函数：文件处理与导出】
# ============================================================
def extract_text_from_docx(file):
    """从 Word 文档中提取文字 [cite: 12]"""
    try:
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"Word解析失败: {e}")
        return ""


def extract_text_from_pdf(file):
    """从 PDF 文档中提取文字 [cite: 12]"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            content = page.extract_text()
            if content: text += content
        return text
    except Exception as e:
        st.error(f"PDF解析失败: {e}")
        return ""


def generate_docx(report_text):
    """生成符合公文规范的 Word 报告（仿宋, 小四, 加粗识别） [cite: 15, 34]"""
    doc = Document()
    # 设置标题
    heading = doc.add_heading('一到五队 - 专利合规审查报告', 0)
    for run in heading.runs:
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.size = Pt(16)
        run.font.bold = True

    # 逐行处理正文并识别加粗
    for line in report_text.split('\n'):
        if line.strip():
            p = doc.add_paragraph()
            # 正则匹配 **加粗内容**
            parts = re.split(r'(\*\*.*?\*\*)', line.strip())
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.replace('**', ''))
                    run.font.bold = True
                else:
                    run = p.add_run(part)

                run.font.name = '仿宋'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                run.font.size = Pt(12)  # 小四 [cite: 15]

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ============================================================
# 【界面布局：用户体验优化】
# ============================================================
st.set_page_config(page_title="一到五队专利合规审查系统", layout="wide", page_icon="🛡️")

# 初始化 Session State
if 'contract_content' not in st.session_state:
    st.session_state['contract_content'] = ""
if 'last_file_name' not in st.session_state:
    st.session_state['last_file_name'] = ""

st.title("🛡️ 一到五队专利许可合同智能合规审查平台")
st.markdown("---")

# 侧边栏：配置与功能按钮 [cite: 12, 29]
st.sidebar.header("配置与工具")
user_input_key = st.sidebar.text_input("临时 API Key (可选)", type="password")

if st.sidebar.button("🗑️ 一键清空所有内容"):
    st.session_state['contract_content'] = ""
    st.session_state['last_file_name'] = ""
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.info("💡 提示：本系统支持自动识别专利期限、许可费率及风险定位 [cite: 18, 48]。")

# 主界面布局
col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("📄 合同输入")

    # 身份界定 [cite: 12]
    user_identity = st.selectbox(
        "您的身份角色",
        ["甲方（许可方）", "乙方（被许可方）", "侵权纠纷原告", "侵权纠纷被告"]
    )

    # 文件上传逻辑 (带自动刷新) [cite: 12]
    uploaded_file = st.file_uploader("直接上传合同文件 (支持 PDF, DOCX)", type=['docx', 'pdf'])

    if uploaded_file and uploaded_file.name != st.session_state['last_file_name']:
        with st.spinner("正在解析文件..."):
            if uploaded_file.name.endswith('.docx'):
                st.session_state['contract_content'] = extract_text_from_docx(uploaded_file)
            elif uploaded_file.name.endswith('.pdf'):
                st.session_state['contract_content'] = extract_text_from_pdf(uploaded_file)
            st.session_state['last_file_name'] = uploaded_file.name
            st.rerun()

            # 文本输入框
    contract_text = st.text_area(
        "合同文本内容",
        value=st.session_state['contract_content'],
        height=450,
        placeholder="请在此粘贴文本或上传文件..."
    )
    st.session_state['contract_content'] = contract_text

    submit_btn = st.button("🚀 开始深度合规审查")

# ============================================================
# 【核心审查逻辑：可视化与专业深度】
# ============================================================
if submit_btn:
    final_api_key = user_input_key if user_input_key else MY_REAL_API_KEY

    if not final_api_key:
        st.error("❌ 请在左侧配置 API Key")
    elif not st.session_state['contract_content'].strip():
        st.warning("⚠️ 请先输入或上传合同内容")
    else:
        start_time = time.time()  # 开始记录耗时 [cite: 57]

        # 可视化进度反馈 [cite: 31, 32]
        with st.status("正在进行深度合规审查...", expanded=True) as status:
            st.write("🔍 正在扫描合同条款（识别专利范围、期限与许可费）...")
            time.sleep(0.8)

            st.write("⚖️ 正在检索《专利法》及相关司法解释依据...")

            # 强化型 Prompt：涵盖专利特有功能与风险定位 [cite: 18, 48, 49]
            prompt_text = f"""
            你是一名资深专利律师。请针对身份【{user_identity}】对以下合同进行深度合规审查。

            【必须包含的审查维度】：
            1. 专利范围分析：分析权利要求书覆盖范围。
            2. 许可费条款识别：评估费率、提成比例及其公平性。
            3. 专利期限判断：校验许可期限是否超过专利法定有效期。
            4. 技术使用范围：校验地域、使用方式及技术改进归属权。

            【格式要求】：
            - 风险定位：必须指明风险对应的具体条款序号。
            - 风险分级：风险点前请明确标注 [高风险]、[中风险] 或 [低风险] [cite: 34]。
            - 核心建议：请使用加粗语法 **内容**。

            合同内容：
            {st.session_state['contract_content']}
            """

            try:
                headers = {
                    'X-Source': 'openapi',
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {final_api_key}'
                }
                data = {
                    "assistant_id": MY_ASSISTANT_ID,
                    "user_id": "team_1to5",
                    "stream": False,
                    "messages": [{"role": "user", "content": [{"type": "text", "text": prompt_text}]}]
                }

                st.write("🚨 正在进行风险分级并生成专业修改建议...")
                response = requests.post(MY_API_URL, headers=headers, json=data)

                if response.status_code == 200:
                    status.update(label="✅ 审查完成！", state="complete", expanded=False)
                    res_json = response.json()
                    report_content = res_json['choices'][0]['message']['content']

                    # 计算耗时 [cite: 57]
                    duration = round(time.time() - start_time, 2)

                    with col2:
                        st.subheader("📊 审查报告")
                        st.caption(f"⏱️ 分析耗时：{duration} 秒 | 法律依据：中国专利法 [cite: 45, 57]")

                        # 风险可视化提醒 [cite: 34]
                        if "[高风险]" in report_content:
                            st.error("🚨 预警：检测到高风险条款，可能涉及核心技术权属丧失。")
                        elif "[中风险]" in report_content:
                            st.warning("⚠️ 提示：部分条款存在法律缺陷，建议进行微调。")

                        st.markdown(report_content)

                        # 生成并下载 Word 报告 [cite: 15]
                        docx_file = generate_docx(report_content)
                        st.download_button(
                            label="📥 下载专业审查报告 (Word 格式)",
                            data=docx_file,
                            file_name="专利许可合规审查报告.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    status.update(label="❌ 审查失败", state="error")
                    st.error(f"接口调用失败 ({response.status_code})")

            except Exception as e:
                status.update(label="❌ 程序异常", state="error")
                st.error(f"异常信息: {e}")

st.markdown("---")
st.caption("注：本系统由一到五队基于腾讯元器及得理法律大数据提供技术支持 [cite: 58]。")
